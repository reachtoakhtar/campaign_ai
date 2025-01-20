import base64
import io
import json
import os
import smtplib
from email.message import EmailMessage
from typing import List, Optional

import requests
from PIL import Image
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langgraph.graph import END, StateGraph, START
from openai import AzureOpenAI
from pydantic import BaseModel, Field

########################
# Fast API configuration
########################
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins or specify your React app's URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.remove(websocket)

    async def send_response(self, data: dict, websocket: WebSocket):
        await websocket.send_json(data)

    async def broadcast(self, message: str):
        for connection in self.active_connections:
            await connection.send_text(message)

manager = ConnectionManager()
#######################

load_dotenv()

# Directory to temporarily store uploaded files
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

vectorstore = Chroma(
    embedding_function=AzureOpenAIEmbeddings(model=os.getenv('TEXT_MODEL'),
    api_key=os.getenv('AZURE_OPENAI_TEXT_API_KEY'),
    azure_endpoint= os.getenv('AZURE_OPENAI_TEXT_ENDPOINT'),
    azure_deployment=os.getenv('AZURE_OPENAI_TEXT_DEPLOYMENT_NAME'),
    openai_api_version=os.getenv('AZURE_OPENAI_API_TEXT_VERSION'),
    ),
    persist_directory="./chroma_db",
)
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,  
    chunk_overlap=50, 
)

class GraphState(BaseModel):
    userPrompt: Optional[str] = None
    generation: Optional[str] = None
    documents: List[str] = []
    image: Optional[str] = None
    imageAnalysis: Optional[str] = None
    imageGeneratefeedbacks: Optional[str] = None
    imageGenerationNum: int = 0
    mailsubject: Optional[str] = None
    mailcontent: Optional[str] = None
    
class GradeAnalysis(BaseModel):
    grade: str = Field(
        description="It gives yes / no if the image aligns with prompt."
    )
    reason: str = Field(
        description="Reason on why the image is not up to the mark."
    )
    
class MailParams(BaseModel):
    subject: str = Field(
        description="Subject for mail."
    )
    content: str = Field(
        description="Content body for mail"
    )

client = AzureOpenAI(api_key=os.getenv('AZURE_OPENAI_DALLE_API_KEY'),api_version=os.getenv('AZURE_OPENAI_DALLE_API_VERSION'),
                     azure_endpoint=os.getenv('AZURE_OPENAI_DALLE_ENDPOINT'))

MAX_GENERATIONS = 3

@app.websocket("/communicate")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            file = await websocket.receive_bytes()
            data = await websocket.receive_json()
            input_text = data["prompt"]
            filename = data["filename"]
            if file:
                 # Validate file type
                if not filename.lower().endswith('.docx'):
                        raise HTTPException(status_code=400, detail="Only DOCX files are allowed")

                # Save the file to a temporary location
                file_path = os.path.join(UPLOAD_FOLDER, filename)
                with open(file_path, "wb") as temp_file:
                    temp_file.write(file)

            await manager.send_response({'response': 'Connected to server'}, websocket)

            def retriever_node(state: GraphState):

                if file is not None:
                # Extract text from the DOCX file
                    document = Document(file_path)
                    docx_text = "\n".join([paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()])

                    # Split the text into chunks
                    chunks = text_splitter.split_text(docx_text)

                    # Store the chunks in Chroma VectorDB
                    vectorstore.add_texts(chunks, metadatas=[{"filename": filename}] * len(chunks))
                    retriever = vectorstore.as_retriever()
                    new_documents = retriever.invoke(state.userPrompt)
                    new_documents = [d.page_content for d in new_documents]
                    state.documents.extend(new_documents)
                    return {"documents": state.documents}
                else:
                    return {"documents": []}

            system_prompt = """
                You are an assistant for creating prompt for dall-e-3 model to create car image for sale based on the user prompt and retrieved context.
                The retrieved context can be blank and hence in that case remember to use only userPrompt for creating prompt for dall-e-3 model to create car image.
                Additional feedback may be provided for why a previous version of the prompts didn't lead to a valid response. Make sure to utilize that feedback to generate a better prompt for dall-e-3 model.
                Only provide the prompt for attractive image that does not include any text in the image and nothing else!
                """

            human_prompt = """
            userPrompt: {userPrompt}
            
            Here is the additional feedback about previous versions of the prompts:
            {feedback}
    
            Context: 
            {context}
    
            Answer:
            """

            rag_prompt = ChatPromptTemplate.from_messages(
            [
            ("system", system_prompt),
            ("human", human_prompt),
            ]
            )
            llm_engine = AzureChatOpenAI(
            model=os.getenv('MODEL'),
            api_key=os.getenv('AZURE_OPENAI_API_KEY'),
            azure_endpoint= os.getenv('AZURE_OPENAI_ENDPOINT'),
            azure_deployment=os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME'),
            openai_api_version=os.getenv('AZURE_OPENAI_API_VERSION'),
            temperature=0.7
            )
            rag_chain = rag_prompt | llm_engine | StrOutputParser()

            def prompt_generation_node(state: GraphState):
                generation = rag_chain.invoke({
                    "context": "\n\n".join(state.documents),
                    "userPrompt": state.userPrompt,
                    "feedback": state.imageGeneratefeedbacks
                })
                return {"generation": generation}

            async def image_generation_node(state: GraphState):
                welcome = 'Welcome image_generation_node--------------------'
                print(welcome)
                await manager.send_response({'response': welcome}, websocket)

                response = client.images.generate(
                model=os.getenv('IMAGE_MODEL'),
                prompt=state.generation,
                size="1024x1024",
                )

                image_url = response.data[0].url
                await manager.send_response({'response': image_url}, websocket)

                base64_image = process_image_to_base64(response.data[0].url)
                await manager.send_response({'response': base64_image}, websocket)
                return {"image": base64_image, "imageGenerationNum": state.imageGenerationNum + 1}

            async def image_grade_node(state: GraphState):
                question = f"""
                The LLM-generated image is produced based on the image prompt, which is derived from the userPrompt and a set of retrieved facts. There is no human in the loop to access or modify the LLM-generated image.
                Your role is to act as an evaluator, evaluate if the LLM-generated image aligns with the userPrompt. Please provide with 'yes' or 'no' based on the following criteria:
                1. Alignment with the userPrompt:
                    Assess how closely the image matches the userPrompt given to the LLM.
                ***Evaluation guide:
                -> 'yes' means the image perfectly aligns with the userPrompt.
                -> 'no' means the image shows minimal or no alignment with the userPrompt.
                ***
                Be honest and objective in your evaluation. Your evaluation should reflect how well the image meets the criteria overall.
                ***REMEBER i only want 'yes' or 'no' and a reason why...if the image is not at all matching then return 'no' else 'yes' ***
                This is the userPrompt:
                {state.userPrompt}
                """
                print('inside image anaysis')
                output = await image_analysis(question, state.image, websocket)
                return {"imageAnalysis":output}

            system_prompt = """
            You are an AI assistant that gives 'yes' or 'no' along with the reason from the image analysis provided. Dont formulate the reason,
            just give the reason provided in image analysis input.
                
            """
            human_prompt = """
            
            This is the image analysis:
            {imageAnalysis}
               
            """


            image_analysis_prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", system_prompt),
                    ("human", human_prompt),
                ]
            )
            image_analysis_grader = (
                image_analysis_prompt
                | llm_engine.with_structured_output(GradeAnalysis)
            )

            async def image_evaluation_node(state: GraphState):
                welcome = "Welcome to image evaluation node-----------"
                print(welcome)
                await manager.send_response({'response': welcome}, websocket)

                image_analysis_grade = image_analysis_grader.invoke(
                    {"imageAnalysis": state.imageAnalysis}
                )
                global image_stream
                if image_analysis_grade.grade.lower() == 'yes':
                    image_stream=state.image
                    return "useful"
                elif state.imageGenerationNum > MAX_GENERATIONS:
                    image_stream=state.image
                    global mail_sub
                    global mail_cont
                    mail_sub = ''
                    mail_cont = ''
                    return "max_image_generation_reached"
                else:
                    return "not relevant"

            async def image_generation_feedback_node(state: GraphState):
                welcome = "Welcome to image_generation_feedback_node-----------"
                print(welcome)
                await manager.send_response({'response': welcome}, websocket)

                question = f"""
                Your role is to give feedback on the LLM generated image that is in base64 encoded. The LLM generated image(base64 encoded) is aligned with the userPrompt.
                Explain how the LLM generated image could be improved so that it aligns with the userPrompt.
                Only provide your feedback and nothing else!
                This is the userPrompt:
                {state.userPrompt}
                """

                output = await image_analysis(question, state.image, websocket)

                feedback = 'Feedback about the image : {}'.format(
                     output
                )
                print('feedback-----',feedback)
                return {"imageGeneratefeedbacks": feedback}

            system_prompt = """
            Your role is to general subject and content body for mail based on the LLM generated image(base64 encoded) and userPrompt.
            In the content body dont include "please reply to this email or contact us" as it will be a system generated mail.
            Only provide subject and content body for mail and nothing else!
            """

            human_prompt = """
            This is the userPrompt:
            {userPrompt}
            LLM generated image: {image}
            """

            mail_caption_subject_generation_feedback_prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", system_prompt),
                    ("human", human_prompt),
                ]
            )

            mail_caption_subject_generation_feedback_chain = (
                mail_caption_subject_generation_feedback_prompt
                | llm_engine.with_structured_output(MailParams)

            )

            def mail_caption_subject_generation_node(state: GraphState):
                print("Welcome to mail_caption_subject_generation_node-----------")
                mail_params = mail_caption_subject_generation_feedback_chain.invoke({
                    "userPrompt": state.userPrompt,
                    "image": state.image
                })
                global mail_sub
                global mail_cont
                mail_sub = mail_params.subject
                mail_cont = mail_params.content

                return {"mailsubject": mail_params.subject, "mailcontent":mail_params.content}

            pipeline = StateGraph(GraphState)

            pipeline.add_node('retrieval_node', retriever_node)
            pipeline.add_node('prompt_generation_node', prompt_generation_node)
            pipeline.add_node('image_generation_node', image_generation_node)
            pipeline.add_node('image_grade_node', image_grade_node)
            pipeline.add_node('image_generation_feedback_node', image_generation_feedback_node)
            pipeline.add_node('mail_param_generation_node', mail_caption_subject_generation_node)

            pipeline.add_edge(START, 'retrieval_node')
            pipeline.add_edge('retrieval_node', 'prompt_generation_node')
            pipeline.add_edge('prompt_generation_node', 'image_generation_node')
            pipeline.add_edge('image_generation_node', 'image_grade_node')
            pipeline.add_conditional_edges(
                'image_grade_node',
                image_evaluation_node,
                {
                    "useful": 'mail_param_generation_node',
                    "not relevant": 'image_generation_feedback_node',
                    "max_image_generation_reached": END

                }
            )
            pipeline.add_edge('image_generation_feedback_node', 'prompt_generation_node')
            pipeline.add_edge('mail_param_generation_node', END)

            rag_pipeline = pipeline.compile()
            inputs = {"userPrompt": input_text}
            image_prompt=''

            outputs = await rag_pipeline.ainvoke(inputs)

            # for output in outputs:
            #     for key, value in output.items():
            #         print(f"Node: {key}")
            #         print(value)

            if file is not None:
                os.remove(file_path)

            await manager.send_response({"image": image_stream, "mailSubject": mail_sub, "mailContent": mail_cont}, websocket)
            return JSONResponse(content={"image": image_stream, "mailSubject": mail_sub, "mailContent": mail_cont}, status_code=200)
    except WebSocketDisconnect:
        manager.disconnect(websocket)
        response = {"status_code": 500, "detail": f"Something went wrong."}
        await manager.send_response(response, websocket)
    except Exception as e:
        await manager.send_response({"response":str(e)}, websocket)


@app.post("/file-process/")
async def file_process(file: UploadFile = File(...)):
    llm_engine = AzureChatOpenAI(
        model=os.getenv('MODEL'),
        api_key=os.getenv('AZURE_OPENAI_API_KEY'),
        azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
        azure_deployment=os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME'),
        openai_api_version=os.getenv('AZURE_OPENAI_API_VERSION'),
        temperature=0.7
    )

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as temp_file:
        temp_file.write(await file.read())

    file_text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
    )
    document = Document(file_path)
    docx_text = "\n".join([paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()])

    # Split the text into chunks
    chunks = file_text_splitter.split_text(docx_text)
    combined_text = " ".join(chunks)

    # First, identify the product type and relevant categories
    category_messages = [
        {"role": "system", "content": """Analyze the text and identify:
        1. The type of product being described
        2. The most relevant 4-6 main feature categories for this type of product

        Return ONLY a JSON object in this exact format:
        {"product_type": "type", "categories": ["category1", "category2", "category3"]}"""},
        {"role": "user", "content": combined_text}
    ]

    category_response = llm_engine.generate([category_messages])
    # Clean and parse JSON response
    try:
        response_text = category_response.generations[0][0].text.strip()
        # Remove any markdown code block indicators if present
        response_text = response_text.replace('```json', '').replace('```', '').strip()
        product_info = json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {response_text}")
        raise e

    # Now get summaries based on identified categories
    summary_messages = [
        {"role": "system", "content": f"""For this {product_info['product_type']}, provide a brief highlight 
        for each of these categories: {', '.join(product_info['categories'])}

        Return ONLY a JSON object where keys are the categories and values are brief highlights under 10 words each.
        Example format: {{"category1": "highlight1", "category2": "highlight2"}}"""},
        {"role": "user", "content": combined_text}
    ]

    summary_response = llm_engine.generate([summary_messages])
    # Clean and parse JSON response
    try:
        response_text = summary_response.generations[0][0].text.strip()
        # Remove any markdown code block indicators if present
        response_text = response_text.replace('```json', '').replace('```', '').strip()
        feature_summary = json.loads(response_text)
        os.remove(file_path)
        return feature_summary
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {response_text}")
        raise e


@app.post("/send-mail/")
async def send_mail(subject: str = Form(...), body: str = Form(...), image: str = Form(...)):
    try:
        send_email(subject, body, image)
        return JSONResponse(content={"message": "Mail sent successfully"}, status_code=200)
    except Exception as e:
        # Clean up and handle errors
        print(repr(e))
        raise HTTPException(status_code=500, detail=f"Failed to send mail: {str(e)}")

def process_image_to_base64(url):
  response = requests.get(url)
  image_bytes = io.BytesIO(response.content)
  img = Image.open(image_bytes)
  img = img.resize((512, 512), Image.Resampling.LANCZOS)
  jpeg_image = io.BytesIO()
  img.save(jpeg_image, format='JPEG')
  jpeg_image.seek(0)
  base64_string = base64.b64encode(jpeg_image.read()).decode('utf-8')
  return "data:image/jpeg;base64," + base64_string

async def image_analysis(question, encoded_image, websocket):
    # Configuration
    headers = {
        "Content-Type": "application/json",
        "api-key": os.getenv('AZURE_OPENAI_API_KEY'),
    }
    # Payload for the request
    payload = {
      "messages": [
        {
          "role": "system",
          "content": [
            {
              "type": "text",
              "text": "You are an AI assistant that helps people find information."
            }
          ]
        },
        {
          "role": "user",
          "content": [
            {
              "type": "text",
              "text": "\n"
            },
            {
              "type": "image_url",
              "image_url": {
                "url": encoded_image
              }
            },
            {
              "type": "text",
              "text": question
            }
          ]
        }
      ],
      "temperature": 0.7,
      "top_p": 0.95,
      "max_tokens": 800
    }
    ENDPOINT = os.getenv('ENDPOINT')
    # Send request
    try:
        await manager.send_response({'response': 'Sending generated image for analysis.'}, websocket)
        response = requests.post(ENDPOINT, headers=headers, json=payload)
        response.raise_for_status()  # Will raise an HTTPError if the HTTP request returned an unsuccessful status code
    except requests.RequestException as e:
        raise SystemExit(f"Failed to make the request. Error: {e}")

    # Handle the response as needed (e.g., print or process)
    output = response.json()
    await manager.send_response(response.json(), websocket)

    return output['choices'][0]['message']['content']

def send_email(subject, body, image,):
    smtp_server = os.getenv('SMTP_SERVER')
    smtp_port= os.getenv('SMTP_PORT')
    email_from = os.getenv("EMAIL_FROM")
    email_to = os.getenv("EMAIL_TO")
    email_password = os.getenv('EMAIL_PASSWORD')

    print(email_from)
    print(email_to)

    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = email_from
    msg['To'] = email_to
    msg.set_content(body)

    try:
        base64_code = image.split(',')[1]
        img_data = base64_code.encode()
        content = base64.b64decode(img_data)

        with open('image.png', 'wb') as fw:
            fw.write(content)

        with open('image.png', 'rb')  as img_file:
            img_data = img_file.read()
            img_type = 'png'
            img_name = 'image.png'
            msg.add_attachment(img_data, maintype='image', subtype=img_type, filename=img_name)

        with smtplib.SMTP(smtp_server, int(smtp_port)) as server:
            server.starttls()
            server.login(email_from, email_password)
            server.send_message(msg)
        return "Success"
    except Exception as e:
        print(f'Error: {e}')
        return "Error"
